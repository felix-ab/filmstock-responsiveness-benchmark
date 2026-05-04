#!/usr/bin/env python3
"""Build V2 Word writeup for the filmstock responsiveness project."""

from __future__ import annotations

import csv
import os
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


from table_geometry import apply_table_geometry, column_widths_from_weights, section_content_width_dxa


REPO_ROOT = Path(__file__).resolve().parents[1]
ANALYSIS_DIR = Path(os.environ.get("FILMSTOCK_ANALYSIS_DIR", REPO_ROOT / "data" / "analysis"))
STATS_DIR = ANALYSIS_DIR / "stats"
ADV_DIR = STATS_DIR / "advanced_suite"
FIG_DIR = Path(os.environ.get("FILMSTOCK_FIGURE_DIR", REPO_ROOT / "figures" / "paper_figures"))
OUT_DOCX = Path(os.environ.get("FILMSTOCK_OUTPUT_DOCX", REPO_ROOT / "paper" / "filmstock_responsiveness_benchmark.docx"))


BASE_FONT = "Times New Roman"
ACCENT = RGBColor(0, 0, 0)
DARK = RGBColor(0, 0, 0)
MUTED = RGBColor(0, 0, 0)
LIGHT_FILL = "FFFFFF"
HEADER_FILL = "F2F2F2"
SOFT_GRAY = "FFFFFF"
MIN_TABLE_FONT_SIZE = 12

MODEL_LABELS = {
    "chatgpt_image_2": "ChatGPT Image 2",
    "xai_grok_imagine": "Grok Imagine",
}
EFFECT_LABELS = {
    "filmstock_cinestill_minus_portra": "CineStill minus Portra",
    "lighting_warm_minus_cool": "Warm practical minus cool ambient",
    "scan_pushed_minus_clean": "Pushed scan minus clean scan",
}
TERM_LABELS = {
    "light": "Lighting",
    "model by light": "Model × lighting",
    "model": "Model",
    "scan": "Development-scan",
    "film": "Filmstock",
    "film by light": "Filmstock × lighting",
    "model by scan": "Model × scan",
    "model by film": "Model × filmstock",
    "light by scan": "Lighting × scan",
    "film by light by scan": "Filmstock × lighting × scan",
    "film by scan": "Filmstock × scan",
}
FEATURE_LABELS = {
    "rgb_mean_b": "RGB blue mean",
    "rgb_mean_g": "RGB green mean",
    "rgb_sd_b": "RGB blue SD",
    "rgb_sd_g": "RGB green SD",
    "rgb_cov_g_b": "RGB green-blue covariance",
    "rgb_cov_r_b": "RGB red-blue covariance",
    "highlight_lab_l": "Highlight Lab L",
    "edge_luma_mean": "Edge luma mean",
    "luma_mean": "Luma mean",
    "highlight_saturation": "Highlight saturation",
    "highlight_warmth_rgb_r_minus_b": "Highlight warmth R-B",
    "highlight_minus_shadow_lab_b": "Highlight-shadow Lab b",
    "highlight_minus_shadow_warmth": "Highlight-shadow warmth",
    "highlight_lab_b": "Highlight Lab b",
    "local_lab_corr_l_b_mean": "Local Lab L-b correlation",
    "lab_corr_l_b": "Global Lab L-b correlation",
    "halation_ring_red_excess_p95": "Halation red-excess p95",
    "highlight_red_excess": "Highlight red excess",
    "palette_weighted_pairwise_lab_distance": "Weighted palette Lab distance",
    "midtone_warmth_rgb_r_minus_b": "Midtone warmth R-B",
    "local_lab_cov_l_a_sd": "Local Lab L-a covariance SD",
    "rgb_mean_r": "RGB red mean",
    "midtone_lab_b": "Midtone Lab b",
    "lab_sd_a": "Lab a SD",
    "lab_chroma_mean": "Lab chroma mean",
    "halation_background_red_excess_mean": "Background red excess",
    "hue_red_orange_share": "Red-orange hue share",
    "hue_teal_cyan_share": "Teal-cyan hue share",
    "rgb_corr_r_g": "RGB red-green correlation",
    "luma_skewness": "Luma skewness",
    "shadow_rolloff_ratio": "Shadow rolloff ratio",
    "hue_circular_mean_turns": "Circular hue mean",
    "hue_blue_share": "Blue hue share",
    "highlight_compression_ratio": "Highlight compression ratio",
    "rgb_corr_r_b": "RGB red-blue correlation",
    "fft_high_freq_log_power_mean": "High-frequency log power",
    "fft_mid_freq_log_power_mean": "Mid-frequency log power",
    "local_lab_corr_l_a_sd": "Local Lab L-a correlation SD",
    "gradient_mag_p95": "Gradient magnitude p95",
    "gradient_mag_mean": "Gradient magnitude",
    "edge_density": "Edge density",
    "edge_density_canny": "Canny edge density",
    "local_contrast_mean": "Local contrast mean",
    "highpass_texture_sd": "High-pass texture SD",
    "texture_highpass_std": "High-pass texture SD",
    "hue_hist_12bin_09": "Hue histogram bin 9",
    "palette_dominant_lab_chroma": "Dominant-palette Lab chroma",
    "local_lab_corr_a_b_sd": "Local Lab a-b correlation SD",
    "gradient_mag_sd": "Gradient magnitude SD",
    "fft_low_freq_log_power_mean": "Low-frequency log power",
    "laplacian_var": "Laplacian variance",
    "palette_dominant_lab_a": "Dominant-palette Lab a",
    "hue_hist_12bin_01": "Hue histogram bin 1",
    "hue_hist_12bin_10": "Hue bin 10 share",
    "lab_corr_a_b": "Global Lab a-b correlation",
    "luma_p01": "Luma p01",
    "luma_p05": "Luma p05",
    "palette_dominant_lab_b": "Dominant-palette Lab b",
    "hue_bin_11to_10": "Hue bin 11-to-10 ratio",
    "top_two_palette_distance": "Top-two palette distance",
    "palette_top2_lab_distance": "Top-two palette Lab distance",
    "midtone_contrast_p05_p95": "Midtone contrast p05-p95",
    "michelson_contrast_p95_p05": "Michelson contrast p95-p05",
    "effective_palette_colors": "Effective palette colors",
    "palette_effective_clusters_6": "Palette effective clusters",
    "palette_entropy_6": "Palette entropy",
}


def read_csv(path: Path) -> list[dict[str, str]]:
    with path.open(newline="") as f:
        return list(csv.DictReader(f))


def fmt(value: str | float, digits: int = 3) -> str:
    v = float(value)
    if abs(v) < 0.001 and v != 0:
        return "< .001"
    return f"{v:.{digits}f}"


def feature_label(name: str) -> str:
    return FEATURE_LABELS.get(name, name.replace("_", " "))


def p_fmt(value: str | float) -> str:
    v = float(value)
    if v < 0.001:
        return "< .001"
    return f"{v:.3f}"


def p_text(value: str | float) -> str:
    p = p_fmt(value)
    return f"p {p}" if p.startswith("<") else f"p = {p}"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text: str, bold: bool = False, color: RGBColor | None = None, size: float = 8.5, align: str = "left") -> None:
    cell.text = ""
    para = cell.paragraphs[0]
    para.paragraph_format.space_after = Pt(0)
    para.paragraph_format.line_spacing = 1.05
    if align == "center":
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == "right":
        para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    else:
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = para.add_run(text)
    run.font.name = BASE_FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), BASE_FONT)
    run.font.size = Pt(max(size, MIN_TABLE_FONT_SIZE))
    run.font.bold = bold
    run.font.color.rgb = color or DARK
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def clear_paragraph_borders(style) -> None:
    p_pr = style.element.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is not None:
        p_pr.remove(p_bdr)


def force_style_font(style, font_name: str) -> None:
    r_pr = style.element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    for key in ("ascii", "hAnsi", "eastAsia", "cs"):
        r_fonts.set(qn(f"w:{key}"), font_name)
    for key in ("asciiTheme", "hAnsiTheme", "eastAsiaTheme", "cstheme"):
        r_fonts.attrib.pop(qn(f"w:{key}"), None)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    fld_char_1 = OxmlElement("w:fldChar")
    fld_char_1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char_2 = OxmlElement("w:fldChar")
    fld_char_2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_1)
    run._r.append(instr_text)
    run._r.append(fld_char_2)


def style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.85)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = BASE_FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), BASE_FONT)
    force_style_font(normal, BASE_FONT)
    normal.font.size = Pt(12)
    normal.font.color.rgb = DARK
    normal.paragraph_format.line_spacing = 1.0
    normal.paragraph_format.space_after = Pt(6)

    for list_style_name in ["List Bullet", "List Number"]:
        list_style = styles[list_style_name]
        list_style.font.name = BASE_FONT
        list_style._element.rPr.rFonts.set(qn("w:eastAsia"), BASE_FONT)
        force_style_font(list_style, BASE_FONT)
        list_style.font.size = Pt(12)
        list_style.font.color.rgb = DARK

    for name, size, bold, before, after, color in [
        ("Title", 16, True, 0, 8, ACCENT),
        ("Subtitle", 12, False, 0, 8, DARK),
        ("Heading 1", 14, True, 14, 6, DARK),
        ("Heading 2", 12.5, True, 10, 4, DARK),
        ("Heading 3", 12, True, 8, 3, DARK),
    ]:
        style = styles[name]
        style.font.name = BASE_FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), BASE_FONT)
        force_style_font(style, BASE_FONT)
        style.font.size = Pt(size)
        style.font.bold = bold
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        clear_paragraph_borders(style)

    header = section.header
    p = header.paragraphs[0]
    p.text = ""
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("Filmstock Responsiveness Benchmark")
    run.font.name = BASE_FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), BASE_FONT)
    run.font.size = Pt(10)
    run.font.color.rgb = DARK
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    footer = section.footer
    p = footer.paragraphs[0]
    p.text = ""
    add_page_number(p)


def add_para(doc: Document, text: str, bold_prefix: str | None = None) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.keep_together = False
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        r.bold = True
        r.font.name = BASE_FONT
        r._element.rPr.rFonts.set(qn("w:eastAsia"), BASE_FONT)
        r.font.size = Pt(12)
        p.add_run(text[len(bold_prefix):])
    else:
        p.add_run(text)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.left_indent = Inches(0.35)
        p.paragraph_format.first_line_indent = Inches(-0.15)
        p.add_run(item)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.keep_together = True
    run = p.add_run(text)
    run.italic = True
    run.font.name = BASE_FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), BASE_FONT)
    run.font.size = Pt(12)
    run.font.color.rgb = DARK


def add_figure(doc: Document, filename: str, caption: str, width_inches: float = 6.6) -> None:
    image_path = FIG_DIR / filename
    if not image_path.exists():
        raise FileNotFoundError(image_path)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run()
    run.add_picture(str(image_path), width=Inches(width_inches))
    add_caption(doc, caption)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], weights: list[float], font_size: float = 8.0, aligns: list[str] | None = None) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for i, h in enumerate(headers):
        set_cell_shading(hdr.cells[i], HEADER_FILL)
        set_cell_text(hdr.cells[i], h, bold=True, size=font_size, align="center")
    aligns = aligns or ["left"] * len(headers)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value, size=font_size, align=aligns[i] if i < len(aligns) else "left")
    width = section_content_width_dxa(doc.sections[0])
    widths = column_widths_from_weights(weights, width)
    apply_table_geometry(table, widths, table_width_dxa=width, indent_dxa=0, cell_margins_dxa={"top": 95, "bottom": 95, "start": 120, "end": 120})
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_callout(doc: Document, text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    set_cell_shading(cell, LIGHT_FILL)
    set_cell_text(cell, text, bold=False, size=9.5)
    width = section_content_width_dxa(doc.sections[0])
    apply_table_geometry(table, [width], table_width_dxa=width, indent_dxa=0, cell_margins_dxa={"top": 135, "bottom": 135, "start": 180, "end": 180})
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_code_block(doc: Document, text: str) -> None:
    for line in text.strip("\n").split("\n"):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.line_spacing = 1.0
        run = p.add_run(line)
        run.font.name = BASE_FONT
        run._element.rPr.rFonts.set(qn("w:eastAsia"), BASE_FONT)
        run.font.size = Pt(12)


def add_math_run(paragraph, text: str, *, italic: bool = False, sub: bool = False, sup: bool = False, size: float = 15) -> None:
    run = paragraph.add_run(text)
    run.font.name = BASE_FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), BASE_FONT)
    run.font.size = Pt(size)
    run.font.color.rgb = DARK
    run.font.italic = italic
    run.font.subscript = sub
    run.font.superscript = sup


def add_display_math(doc: Document, parts: list[tuple[str, str | None]], *, size: float = 15) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_together = True
    for text, style in parts:
        flags = set((style or "").split())
        add_math_run(
            p,
            text,
            italic="italic" in flags,
            sub="sub" in flags,
            sup="sup" in flags,
            size=size,
        )


EQUATION_PARTS = {
    "eq_standardization.png": [
        ("z", "italic"), ("ij", "sub"), (" = (", None), ("x", "italic"), ("ij", "sub"),
        (" − ", None), ("x̄", "italic"), ("j", "sub"), (") / ", None), ("s", "italic"), ("j", "sub"),
    ],
    "eq_distance.png": [
        ("D", "italic"), ("i", "sub"), (" = [Σ", None), ("j∈S", "sub"), ("(", None),
        ("z", "italic"), ("high, ij", "sub"), (" − ", None), ("z", "italic"), ("low, ij", "sub"),
        (")", None), ("2", "sup"), ("]", None), ("1/2", "sup"),
    ],
    "eq_delta.png": [
        ("Δ", None), ("i", "sub"), (" = ", None), ("D", "italic"), ("i, Grok", "sub"),
        (" − ", None), ("D", "italic"), ("i, ChatGPT", "sub"),
    ],
    "eq_pca_covariance.png": [
        ("C", "italic"), (" = ", None), ("Z", "italic"), ("T", "sup"), ("Z", "italic"),
        (" / (", None), ("n", "italic"), (" − 1)", None),
    ],
    "eq_pca_eigenvector.png": [
        ("C", "italic"), ("v", "italic"), ("k", "sub"), (" = ", None), ("λ", None),
        ("k", "sub"), ("v", "italic"), ("k", "sub"),
    ],
    "eq_pca_score.png": [
        ("t", "italic"), ("ik", "sub"), (" = ", None), ("z", "italic"), ("i", "sub"),
        (" · ", None), ("v", "italic"), ("k", "sub"),
    ],
    "eq_wls_weights.png": [
        ("w", "italic"), ("i", "sub"), (" ≈ 1 / ", None), ("σ̂", None), ("i", "sub"),
        ("2", "sup"),
    ],
    "eq_appendix_standardization.png": [
        ("z", "italic"), ("ij", "sub"), (" = (", None), ("x", "italic"), ("ij", "sub"),
        (" − ", None), ("x̄", "italic"), ("j", "sub"), (") / ", None), ("s", "italic"), ("j", "sub"),
    ],
    "eq_appendix_distance.png": [
        ("D", "italic"), ("i", "sub"), (" = [Σ", None), ("j∈S", "sub"), ("(", None),
        ("z", "italic"), ("high, ij", "sub"), (" − ", None), ("z", "italic"), ("low, ij", "sub"),
        (")", None), ("2", "sup"), ("]", None), ("1/2", "sup"),
    ],
}


def add_equation(doc: Document, filename: str, latex: str, width_inches: float | None = None) -> None:
    parts = EQUATION_PARTS.get(filename)
    if parts is None:
        parts = [(latex, None)]
    add_display_math(doc, parts, size=15)


def add_numbered_item(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Inches(0.35)
    p.paragraph_format.first_line_indent = Inches(-0.15)
    p.add_run(text)


def primary_distance_rows() -> list[list[str]]:
    rows = read_csv(STATS_DIR / "paired_model_distance_tests.csv")
    out = []
    for effect in EFFECT_LABELS:
        r = next(row for row in rows if row["effect_type"] == effect and row["distance_metric"] == "distance_all_selected_features")
        out.append([
            r["effect_label"],
            fmt(r["mean_chatgpt"]),
            fmt(r["mean_xai"]),
            fmt(r["mean_difference_xai_minus_chatgpt"]),
            f"[{fmt(r['ci95_low'])}, {fmt(r['ci95_high'])}]",
            p_fmt(r["paired_t_p_value"]),
        ])
    return out


def bootstrap_rows() -> list[list[str]]:
    rows = read_csv(ADV_DIR / "hierarchical_bootstrap_distance_tests.csv")
    out = []
    for effect in EFFECT_LABELS:
        r = next(row for row in rows if row["effect_type"] == effect and row["distance_metric"] == "distance_all_selected_features")
        out.append([
            r["effect_label"],
            fmt(r["mean_difference_xai_minus_chatgpt"]),
            f"[{fmt(r['hierarchical_boot_ci95_low'])}, {fmt(r['hierarchical_boot_ci95_high'])}]",
            p_fmt(r["hierarchical_boot_p_two_sided"]),
        ])
    return out


def permanova_rows() -> list[list[str]]:
    rows = read_csv(ADV_DIR / "blocked_permanova_pc10.csv")
    rows.sort(key=lambda r: float(r["partial_r2"]), reverse=True)
    return [[TERM_LABELS.get(r["term_label"], r["term_label"]), fmt(r["partial_r2"]), fmt(r["pseudo_f"]), p_fmt(r["permutation_p_value"]), p_fmt(r["bh_fdr_p"])] for r in rows[:8]]


def manova_rows() -> list[list[str]]:
    rows = read_csv(STATS_DIR / "manova_pillai_permutation.csv")
    rows.sort(key=lambda r: float(r["pillai_trace"]), reverse=True)
    return [[TERM_LABELS.get(r["term_label"], r["term_label"]), fmt(r["pillai_trace"]), p_fmt(r["permutation_p_value"]), p_fmt(r["bh_fdr_p"])] for r in rows[:8]]


def regression_rows() -> list[list[str]]:
    rows = read_csv(STATS_DIR / "regression_distance_coefficients.csv")
    terms = ["xAI model", "Lighting effect", "Scan effect", "xAI by lighting effect", "xAI by scan effect", "Run block 2", "Run block 3"]
    display_terms = {
        "xAI model": "Grok Imagine indicator",
        "Lighting effect": "Lighting contrast",
        "Scan effect": "Pushed-scan contrast",
        "xAI by lighting effect": "Grok × lighting contrast",
        "xAI by scan effect": "Grok × pushed-scan contrast",
    }
    out = []
    for term in terms:
        vals = {}
        for method in ["OLS HC3", "Two-stage WLS", "Huber IRLS"]:
            r = next(row for row in rows if row["metric"] == "distance_all_selected_features" and row["method"] == method and row["term"] == term)
            vals[method] = f"{fmt(r['estimate'])} ({p_text(r['p_value'])})"
        out.append([display_terms.get(term, term), vals["OLS HC3"], vals["Two-stage WLS"], vals["Huber IRLS"]])
    return out


def pls_rows() -> list[list[str]]:
    rows = read_csv(ADV_DIR / "pls_da_leave_run_out_cv.csv")
    out = []
    for target in ["model_key", "light_key", "film_key", "scan_key"]:
        best = max([r for r in rows if r["target"] == target], key=lambda r: float(r["leave_run_out_auc"]))
        out.append([best["target_label"], best["n_components"], fmt(best["leave_run_out_auc"]), fmt(best["leave_run_out_accuracy"]), fmt(best["leave_run_out_rmse"])])
    return out


def regularized_rows() -> list[list[str]]:
    rows = read_csv(ADV_DIR / "regularized_leave_run_out_cv.csv")
    rows.sort(key=lambda r: float(r["leave_run_out_auc"]), reverse=True)
    return [[r["target_label"], r["method"], fmt(r["leave_run_out_auc"]), fmt(r["leave_run_out_accuracy"]), r["selected_c_median"]] for r in rows]


def variance_component_rows() -> list[list[str]]:
    rows = read_csv(ADV_DIR / "image_feature_variance_components.csv")
    out = []
    for model in ["chatgpt_image_2", "xai_grok_imagine"]:
        subset = [r for r in rows if r["model_key"] == model]
        def avg(key: str) -> str:
            return fmt(sum(float(r[key]) for r in subset) / len(subset))
        out.append([MODEL_LABELS[model], avg("condition_share"), avg("run_block_share"), avg("condition_by_run_share"), avg("within_run_replicate_share")])
    return out


def pca_loading_rows() -> list[list[str]]:
    rows = read_csv(ANALYSIS_DIR / "pca_loadings.csv")
    out = []
    for component in ["PC1", "PC2"]:
        subset = [r for r in rows if r["component"] == component]
        positive = sorted(subset, key=lambda r: float(r["loading"]), reverse=True)[:4]
        negative = sorted(subset, key=lambda r: float(r["loading"]))[:4]
        pos_text = "; ".join(f"{feature_label(r['feature'])} ({float(r['loading']):+.3f})" for r in positive)
        neg_text = "; ".join(f"{feature_label(r['feature'])} ({float(r['loading']):+.3f})" for r in negative)
        out.append([component, pos_text, neg_text])
    return out


def top_screening_rows(effect: str, n: int = 8) -> list[list[str]]:
    rows = [r for r in read_csv(STATS_DIR / "all_feature_model_screening.csv") if r["effect_type"] == effect]
    return [[feature_label(r["feature"]), r["feature_family"].replace("_", " "), fmt(r["abs_standardized_model_gap"]), p_fmt(r["bh_fdr_p_all_features"])] for r in rows[:n]]


def top_reliability_rows(n: int = 10) -> list[list[str]]:
    rows = read_csv(ADV_DIR / "cross_run_feature_reliability.csv")
    return [[MODEL_LABELS.get(r["model_key"], r["model_key"]), r["effect_label"], r["feature"], fmt(r["standardized_mean_run_effect"]), f"{r['same_sign_run_count']}/3"] for r in rows[:n]]


def inventory_title(row: dict[str, str]) -> str:
    title = row["title"].replace("_", " ").replace(".png", "")
    parts = [part for part in title.split() if part.lower() not in {"v2", "existing"} and not part.isdigit()]
    clean = " ".join(parts)
    replacements = {
        "pca by model": "PCA by model",
        "pca by lighting": "PCA by lighting",
        "pca by filmstock": "PCA by filmstock",
        "pca by run block": "PCA by run block",
        "blocked permanova": "Blocked PERMANOVA",
        "manova vs permanova": "MANOVA and PERMANOVA comparison",
        "responsiveness distances": "Responsiveness distances",
        "distance distributions": "Distance distributions",
        "paired lighting lines": "Paired lighting lines",
        "hierarchical bootstrap forest": "Hierarchical bootstrap forest",
        "curated targeted feature heatmap": "Curated targeted feature heatmap",
        "all feature screening volcano": "All-feature screening volcano",
        "top lighting screening features": "Top lighting screening features",
        "top filmstock screening features": "Top filmstock screening features",
        "pls vip features": "PLS VIP features",
        "regularized coefficients": "Regularized coefficients",
        "feature variance components": "Feature variance components",
        "seed to seed variation": "Seed-to-seed variation",
        "run block variation": "Run-block variation",
        "cross run reliable features": "Cross-run reliable features",
        "robust outlier diagnostics": "Robust outlier diagnostics",
        "selected feature correlation heatmap": "Selected feature correlation heatmap",
        "corner store contact sheet": "Corner-store contact sheet",
        "apartment contact sheet": "Apartment contact sheet",
        "backstage contact sheet": "Backstage contact sheet",
        "model difference forest": "Model-difference forest plot",
        "interaction heatmap": "Interaction heatmap",
        "nonparametric ordering": "Nonparametric ordering plot",
    }
    return replacements.get(clean.lower(), clean.capitalize())


def inventory_role(row: dict[str, str]) -> str:
    role = row["role"]
    if role == "Copied existing diagnostic":
        return "Supplementary diagnostic"
    return role


def inventory_rationale(row: dict[str, str]) -> str:
    rationale = row["rationale"]
    replacements = {
        "Useful secondary evidence; less visually separated than lighting.": "Secondary evidence; less visually separated than lighting.",
        "Visual proof-of-concept for the prompt manipulations.": "Visual reference for the prompt manipulations.",
        "Shows spread and outliers; useful but busy.": "Shows spread and outliers; retained as a secondary distribution view.",
        "Important audit trail, but visually dense.": "Important audit trail; retained as a dense supplemental figure.",
        "Good robustness comparison, but somewhat redundant with PERMANOVA.": "Secondary robustness comparison to the PERMANOVA result.",
        "Useful for technical; better as appendix.": "Technical shrinkage-model detail retained in the appendix.",
        "Useful but technical; better as appendix.": "Technical shrinkage-model detail retained in the appendix.",
        "Useful supplement to seed variation.": "Supplemental replication view for run-block variation.",
        "Strong but label-heavy; keep for appendix.": "Strong reliability evidence retained in the appendix.",
        "Useful diagnostic but not core narrative.": "Diagnostic view retained outside the main narrative.",
        "Useful but dense feature-correlation diagnostic with top labels for readability.": "Dense feature-correlation diagnostic with top labels for readability.",
    }
    return replacements.get(rationale, rationale)


def build_doc() -> None:
    doc = Document()
    style_document(doc)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run("Felix Brener\nSTT217, Spring 2026\nMay 1, 2026")
    r.font.name = BASE_FONT
    r._element.rPr.rFonts.set(qn("w:eastAsia"), BASE_FONT)
    r.font.size = Pt(12)
    r.font.color.rgb = DARK

    title = doc.add_paragraph(style="Title")
    title.add_run("Measuring Filmstock and Lighting Prompt Responsiveness in Image Generation Models")

    add_heading(doc, "Abstract", 1)
    add_para(doc, "This study proposes a filmstock-based benchmark for photographic prompt responsiveness in image generation models. Instead of asking only whether a model follows literal prompt content, the study asks whether responsiveness to filmstock, lighting, and development-scan language can be measured statistically, and whether models differ in how they respond to those artistic controls. The experiment compares ChatGPT Image 2 and Grok Imagine using a balanced factorial design with 288 generated images: two models, three base scenes, two film stocks, two lighting conditions, two development-scan conditions, two within-run replicates, and three independent generation runs. Each image was converted into 195 quantitative features describing color, Lab covariance, hue distribution, split-toning, tonal curve shape, grain and texture, Fourier frequency structure, palette organization, edge structure, and halation proxies.")
    add_para(doc, "The empirical result is comparative rather than a single model ranking: photographic prompt responsiveness is multidimensional. Lighting is the strongest and most reliable prompt factor, filmstock and scan are detectable but subtler, and the most robust model difference appears as a model-by-lighting interaction. After reducing the 195 extracted image features to a lower-dimensional multivariate representation, blocked PERMANOVA found significant effects for lighting, model-by-lighting interaction, model, scan, filmstock, and filmstock-by-lighting interaction. Hierarchical bootstrap inference confirmed a significant model difference for warm-practical lighting responsiveness but not for the filmstock or scan contrasts. The conclusion is that a filmmaker's benchmark should evaluate not only correctness, but the controllability of aesthetic dimensions that matter in professional creative workflows.")

    add_heading(doc, "1. Filmstock Prompt Responsiveness as a Benchmarking Problem", 1)
    add_para(doc, "Image generation benchmarks face a practical creative problem. In applied image-generation workflows, a model can be technically impressive while still offering weak aesthetic control. It may render a plausible person, a readable object, and a convincing scene, yet respond only faintly to the language that photographers, cinematographers, designers, and art directors use to shape a specific look. A prompt like “a man in a red raincoat riding a blue bicycle past a yellow taxi on a city street” tests a model's ability to place objects and attributes. A prompt that adds “CineStill 800T,” “warm tungsten practical lights,” or “pushed one stop in development” tests a different capability: whether the model understands the visual consequences of a production-style photographic instruction. For that reason, the study is intentionally about photographic prompt responsiveness rather than ordinary prompt adherence.")
    add_para(doc, "In practical discussions about image-generation tools, models are often compared through informal judgments about taste, controllability, and whether the system can be steered toward a specific visual result. Those impressions matter, but they are rarely translated into a reproducible statistical design. This study turns one part of that intuition into a controlled experiment: when a photographic cue is added to a prompt, the analysis measures whether the image changes in the expected feature dimensions.")
    add_para(doc, "As of May 2026, text-to-image benchmarks cover several important evaluation lenses. DrawBench was introduced with Imagen to test text-to-image prompt following and human preference across challenging prompts (Saharia et al., 2022). PartiPrompts provides a large prompt set for evaluating content-rich generation (Yu et al., 2022). GenEval focuses on object-level alignment, including count, position, color, and object co-occurrence (Ghosh et al., 2023). T2I-CompBench++ targets compositional skills such as attribute binding, spatial relationships, and complex compositions (Huang et al., 2023). TIFA, VQAScore, and GenAI-Bench use visual question answering, image-to-text evaluation, or compositional text-to-visual testing to evaluate whether an image is faithful to the text (Hu et al., 2023; Li et al., 2024; Lin et al., 2024). HEIM takes a broader holistic-evaluation view across multiple aspects of text-to-image behavior (Lee et al., 2023). HPSv2 and related preference models ask which images humans tend to prefer (Wu et al., 2023). Recent color-focused work such as GenColorBench and ColorConceptBench evaluates more specialized color controllability and color-concept understanding (Butt et al., 2025; Ruan et al., 2026). DALL-E 3's paper emphasizes improved prompt following through better captions (Betker et al., 2023). These are important evaluation lenses, but none are primarily designed to ask whether a model responds to filmstock language, lighting design, scan treatment, or other professional aesthetic controls.")
    add_para(doc, "Against that benchmark landscape, the extension proposed here is deliberately narrower and more production-oriented: treat filmstock, lighting, and development-scan terms as experimental factors, then test whether those factors move generated images through measurable color, tone, texture, and covariance dimensions. This study therefore treats filmstock responsiveness as a small-scale benchmark prototype for a complementary axis: a filmmaker's benchmark. The point is not to replace object-adherence benchmarks. A useful creative model should still render anatomy, text, objects, spatial relationships, and prompt details correctly. The point is that, for creative professionals, correctness is not enough. A controllable image model should also move in the right visual dimension when given an artistic instruction.")
    add_para(doc, "The central estimand is intentionally simple: the matched feature-space displacement caused by changing one photographic prompt factor while holding the rest of the design fixed. The additional multivariate, regression, robust-estimation, nonparametric, and predictive analyses are used to test whether that displacement is reliable, interpretable, and robust to alternative statistical views, not to turn the project into a model leaderboard.")
    add_para(doc, "The artistic language was narrowed to filmstock, lighting, and development-scan conditions because these terms should leave measurable traces in color space, local color covariance, tonal distribution, contrast, grain-like texture, and halation proxies. Other artistic terminology, such as camera angle, lens rendering, staging density, blocking, production design, or editorial mood, is important but harder to operationalize with a small feature set. The 288 images in this study should therefore be read as a deliberately small experimental sample, not a finished industry-scale benchmark. A larger benchmark would expand to thousands of photographic and cinematic terms, many more scenes, reference photographs, and human expert ratings so that the evaluation becomes broad enough to inform practical creative-tool assessment.")

    doc.add_page_break()
    add_heading(doc, "2. Study Design", 1)
    add_para(doc, "The study used a balanced factorial design. The base scenes were held constant while three binary stylistic factors were manipulated: film stock, lighting, and development-scan condition. This kept the experiment thematically coherent and avoided uncontrolled variation from mixing unrelated style terms.")
    add_table(
        doc,
        ["Factor", "Levels", "Experimental function"],
        [
            ["Model", "ChatGPT Image 2; Grok Imagine", "Primary model comparison"],
            ["Scene", "Apartment editorial; backstage musician; corner store fashion", "Repeated base-prompt block"],
            ["Film stock", "Kodak Portra 400; CineStill 800T", "Film color, palette, and halation factor"],
            ["Lighting", "Cool ambient; warm tungsten practical", "Color temperature, split-tone, and highlight factor"],
            ["Development-scan", "Clean box-speed scan; pushed one-stop scan", "Contrast, shadow, and grain factor"],
            ["Within-run replicate", "Two images per exact condition", "Seed-to-seed variation within a run"],
            ["Run block", "Three complete generation batches", "Independent repeated blocks of the full design"],
            ["Total", "2 × 3 × 2 × 2 × 2 × 2 × 3 = 288", "Balanced full factorial sample"],
        ],
        [1.25, 2.25, 2.5],
        font_size=7.8,
    )
    add_para(doc, "The prompt set was chosen to be specific enough for professional photographic interpretation but controlled enough for statistical comparison. The three base scenes resemble editorial-style image workflows: a 25-year-old woman in a lived-in city apartment, a 22-year-old man in a backstage green room, and a 30-year-old woman outside a corner store. These scenes were intentionally human-centered because skin tone, mixed light, shadow detail, and wardrobe/background color give the filmstock and lighting factors room to appear. At the same time, all three prompts use the same camera framing assumptions, which keeps the base content from overwhelming the stylistic manipulations.")
    add_para(doc, "The factor choices were deliberately photographic rather than decorative. Kodak Portra 400 and CineStill 800T are familiar 35mm color-negative references with different expected visual signatures: Portra suggests natural skin tone, softer pastel color response, and gentle highlight rolloff, while CineStill 800T suggests tungsten balance, stronger cinematic color separation, and possible red halation around bright highlights. The lighting factor contrasts cool diffuse ambient light with warm tungsten practical light plus subtle neon spill, which should be visible in Lab b, RGB warmth, highlight-shadow color separation, and local lightness-color covariance. The scan/development factor contrasts a clean box-speed Frontier-style scan with a pushed one-stop scan, which should affect lower-tail luminance, contrast, shadow density, and grain-like texture.")
    add_para(doc, "These are not arbitrary aesthetic labels. They are production-facing descriptors a photographer, colorist, or art director could plausibly use when specifying the desired behavior of an image. That makes the design relevant to professional use while also making it statistically tractable: each factor has a concrete expected visual direction and can be evaluated through color, tone, texture, and covariance features.")
    add_para(doc, "The prompt language avoided ranges, slashes, and multiple-choice phrasing because those forms add ambiguity for image models. Every prompt used concrete photographic settings: square photorealistic editorial 35mm photograph, natural proportions, realistic skin texture, eye-level camera, 50mm lens, medium shot, no text, no logos, no watermark, and no artificial digital smoothness.")
    add_para(doc, "Closed cloud models limit experimental control over seeds and samplers, but they are also commercially relevant systems used in practical creative workflows. The three complete generation runs make it possible to evaluate whether effects repeat reliably across batches despite the hidden sampler details. In the statistical models, run block is treated as a blocking factor, and within-run replicate is used to estimate seed-to-seed variation under exact prompt conditions.")
    add_figure(doc, "v2_26_contact_sheet_corner_store.png", "Figure 1. Corner-store contact sheet from one run. The image grid gives a concrete visual reference for the filmstock and lighting manipulations.", width_inches=6.7)

    add_heading(doc, "3. Image Feature Extraction", 1)
    add_para(doc, "Each generated PNG was converted into a row of numeric image measurements. The resulting analytic matrix contains 288 images and 195 image-derived variables. The raw images were 1024 × 1024 RGB files, but the statistical analysis does not use raw pixels directly. It uses summary features chosen to represent the kinds of measurable visual shifts that filmstock, lighting, and development-scan language should plausibly create.")
    add_para(doc, "The first feature group describes global color and perceptual color balance: RGB means and standard deviations, Lab lightness and a/b channels, saturation, hue entropy, and warm-color shares. These variables are especially relevant for distinguishing cool ambient light from warm tungsten practical light and for detecting whether filmstock language changes the overall palette.")
    add_para(doc, "The second group describes local color structure. The image is divided into patches, and covariance or correlation is computed among Lab lightness, Lab a, and Lab b values. These local covariance features matter because photographic style is not only average color; it is also how color and brightness travel together across highlights, shadows, faces, clothing, and background regions.")
    add_para(doc, "A third group focuses on split-toning and tonal behavior. Separate color summaries are computed for shadow, midtone, and highlight regions, along with highlight-minus-shadow differences, lower-tail luminance, highlight rolloff, shadow compression, skewness, and kurtosis. This feature family is directly relevant to warm practical lighting and pushed development, where the expected signal may appear in shadow density or highlight color rather than in the whole-image mean.")
    add_para(doc, "The remaining feature groups measure texture, spatial structure, palette organization, and halation proxies. High-pass texture, Laplacian variance, Fourier frequency power, Canny edge density, gradient magnitude, six-cluster Lab palette entropy, palette distances, and red-excess measures near bright regions were included to capture grain-like detail, edge complexity, color organization, and CineStill-like halation.")
    add_para(doc, "Each numeric feature was standardized before analysis, and each matched prompt contrast was measured as a Euclidean distance in standardized feature space. A filmstock distance compares CineStill against Portra while holding model, run, scene, lighting, scan, and replicate fixed. Lighting and scan distances are constructed the same way. These distances are unitless because they are computed after z-scoring the extracted image features; a larger value means the manipulated prompt factor produced more multivariate movement in the selected feature set.")
    add_para(doc, "Quality control confirmed that all 288 generated files were readable 1024 × 1024 RGB images. No feature row contained missing values after extraction, so the statistical analysis required no imputation or case deletion.")

    add_heading(doc, "4. Statistical Methods", 1)
    add_para(doc, "The analysis treats the study as a blocked factorial experiment with high-dimensional image-derived outcomes. There are two related data objects. The first is the image-feature matrix X, with 288 rows and 195 standardized feature columns. The second is a matched-contrast distance table, where each row measures how far an image moves when one prompt factor changes while the other experimental factors are held fixed.")
    add_para(doc, "The results use three related measurements. A standardized feature value is a z-score for one image and one extracted feature. A feature effect is a high-minus-low prompt contrast for one feature, often reported after dividing by that feature's sample standard deviation. A responsiveness distance is different: it is the Euclidean length of many z-score differences across a selected feature set. In short, feature effects identify which visual measurements changed, while responsiveness distances summarize how far an image moved across many measurements at once.")
    add_para(doc, "The analysis hierarchy is as follows. The primary response variable is the matched all-feature responsiveness distance. PCA, MANOVA, and PERMANOVA provide multivariate confirmation that the factorial prompt conditions move the broader feature space. Targeted Lab, split-tone, halation, texture, and palette features explain the photographic meaning of those movements. Regression, WLS, Huber IRLS, nonparametric tests, PLS, LASSO/ridge, variance shares, and outlier diagnostics are treated as robustness checks or exploratory diagnostic views rather than separate headline claims.")

    doc.add_page_break()
    add_heading(doc, "4.1 Standardization and Response Distances", 2)
    add_para(doc, "Because the extracted variables have different natural scales, every feature was standardized before modeling. For image i and feature j, the standardized feature value was computed as")
    add_equation(doc, "eq_standardization.png", r"z_{ij}=\frac{x_{ij}-\bar{x}_{j}}{s_{j}}")
    add_para(doc, "where the numerator centers each observation by the feature-specific sample mean and the denominator rescales by the feature-specific sample standard deviation. Prompt responsiveness was then defined as a Euclidean distance in standardized feature space:")
    add_equation(doc, "eq_distance.png", r"D_i=\sqrt{\sum_{j\in S}\left(z_{\mathrm{high},ij}-z_{\mathrm{low},ij}\right)^2}")
    add_para(doc, "Here S is the selected feature set: either all 195 features or a feature family such as color, structure, or texture. The value of Dᵢ is not measured in pixels or raw color-channel units; it is the distance between two images after the relevant features have been converted to z-scores.")
    add_para(doc, "This construction makes the experiment paired by design. For example, a lighting distance compares warm practical lighting against cool ambient lighting while holding model, run block, scene, filmstock, scan condition, and replicate fixed. This pairing removes a large amount of nuisance variation and focuses the response variable on the manipulated prompt factor.")

    add_heading(doc, "4.2 Paired Inference and Hierarchical Bootstrap", 2)
    add_para(doc, "The primary model comparison uses paired differences for each matched condition:")
    add_equation(doc, "eq_delta.png", r"\Delta_i=D_{i,\mathrm{Grok}}-D_{i,\mathrm{ChatGPT}}")
    add_para(doc, "Ordinary paired t tests were computed for filmstock, lighting, and scan responsiveness. Because the three generation runs are repeated batches rather than independent pixels, the analysis also used a hierarchical bootstrap. The bootstrap resamples run blocks first, then resamples matched contrasts within the selected runs. This gives confidence intervals and p-values that respect the blocked generation structure.")

    add_heading(doc, "4.3 Principal Components Analysis", 2)
    add_para(doc, "PCA was applied to the 195 standardized image features to summarize the multivariate feature space. The computation starts from the standardized matrix Z, where each column has mean 0 and standard deviation 1. The covariance matrix is")
    add_equation(doc, "eq_pca_covariance.png", r"C=\frac{Z^{T}Z}{n-1}")
    add_para(doc, "Each principal component is an eigenvector of C:")
    add_equation(doc, "eq_pca_eigenvector.png", r"Cv_k=\lambda_k v_k")
    add_para(doc, "Each image's score on component k is the dot product between that image's standardized feature vector and the component loading vector:")
    add_equation(doc, "eq_pca_score.png", r"t_{ik}=z_i\cdot v_k")
    add_para(doc, "A high component score means the image has high values on features with positive loadings and low values on features with negative loadings; a low score means the reverse. The axis labels in the PCA figures were produced by ranking the 195 loadings for each component by absolute value, inspecting the largest positive and negative loadings, and converting those loading patterns into a short photographic description. The first ten principal components explain 75.5 percent of total standardized feature variance, so PC1-PC10 provide a compact but still information-rich representation for later MANOVA and PERMANOVA tests.")
    add_table(
        doc,
        ["Component", "Largest positive loadings: high component score", "Largest negative loadings: low component score"],
        pca_loading_rows(),
        [0.75, 2.7, 2.7],
        font_size=6.9,
        aligns=["center", "left", "left"],
    )

    add_heading(doc, "4.4 Blocked MANOVA and PERMANOVA", 2)
    add_para(doc, "To test whether the factorial prompt conditions moved the multivariate image feature space, PC1 through PC10 were used as the response matrix. These are unitless standardized component scores, not raw pixel values. Blocked MANOVA evaluates multivariate mean separation with Pillai's trace. PERMANOVA evaluates the same design from a distance-based perspective: first compute Euclidean distances among images in PC1-PC10 space, then evaluate whether images sharing a design factor are closer together than would be expected under run-blocked permutations.")
    add_para(doc, "The blocked permutation step is important. Instead of freely shuffling labels across the full 288-image data set, labels are permuted within run block. This preserves the fact that each generation run was a batch replicate. The PERMANOVA partial R² values are therefore interpreted as the share of PC1-PC10 feature-space variation attributable to each term after accounting for the other modeled factors and after respecting the blocked generation structure.")

    add_heading(doc, "4.5 Regression, WLS, and Robust Estimation", 2)
    add_para(doc, "The distance table also supports regression modeling. The main linear model predicts all-feature responsiveness distance from model, prompt factor, model-by-factor interactions, scene, run block, and replicate. OLS with HC3 robust standard errors is used as the baseline because HC3 is less fragile when leverage differs across observations.")
    add_para(doc, "Two-stage weighted least squares is used as a heteroskedasticity check. First, fit the OLS model and compute residuals eᵢ. Second, model the typical residual magnitude as a function of fitted values or design cells. Third, convert the estimated error spread into weights:")
    add_equation(doc, "eq_wls_weights.png", r"w_i\approx\frac{1}{\hat{\sigma}_i^2}")
    add_para(doc, "The WLS fit then gives less influence to observations whose design cells have less precise distance measurements, rather than pretending every matched contrast has equal precision.")
    add_para(doc, "Huber robust regression is used as an additional sensitivity analysis. Huber IRLS begins with a standard fit, computes residuals, gives near-full weight to small residuals, gives reduced weight to large residuals, and refits repeatedly until the coefficients stabilize. This is different from deleting images. A visually unusual output can still contribute to the model, but it cannot dominate the coefficient estimates merely because squared-error loss would otherwise give it very large influence.")

    add_heading(doc, "4.6 Kruskal-Wallis, Friedman ANOVA, and Leave-Run-Out Classification", 2)
    add_para(doc, "Because generated-image features need not satisfy normal-error assumptions, the analysis also used Kruskal-Wallis and Friedman repeated-measures tests to evaluate whether the ordering of filmstock, lighting, and scan distances remains visible without relying on parametric regression assumptions. For predictive validation, PLS-DA, LASSO logistic regression, and ridge logistic regression were fit under leave-run-out validation. These models evaluate whether the extracted features can recover known labels such as model, lighting, filmstock, and scan condition in a held-out generation run.")

    add_heading(doc, "4.7 Feature Variance Shares and Robust Mahalanobis Distances", 2)
    add_para(doc, "Finally, feature-level variance was decomposed into prompt-condition signal, run-block variance, condition-by-run variation, and within-run replicate variation. For each model and each standardized feature, the design was treated as a crossed layout: prompt condition equals scene × filmstock × lighting × scan, run block is the batch factor, condition × run captures batch-specific prompt behavior, and replicate captures within-run seed-to-seed variation. Each component was converted into a share of that feature's total sum of squares, and the figure reports the mean share across all 195 features.")
    add_para(doc, "This addresses the practical concern that closed models may vary from batch to batch. If run-block variance were large, a single generation batch would be a weak basis for evaluating model behavior. Robust Mahalanobis distances in PC1-PC10 space were used to flag unusual generated images without letting those images define the covariance structure being used to detect them.")

    doc.add_page_break()
    add_heading(doc, "5. Results", 1)
    add_heading(doc, "5.1 PC1-PC10 Feature-Space Separation", 2)
    add_para(doc, "The PCA plots show two important facts at once. First, the two models overlap substantially, so the study is not merely detecting that one model has a completely different global look. Second, lighting produces a visibly stronger separation than either model or filmstock alone.")
    add_para(doc, "The interpretation of the PCA axes is reproducible rather than impressionistic. PC1 was labeled by inspecting the largest loadings on its eigenvector: the positive side includes RGB blue mean, RGB green mean, RGB blue/green variation, and highlight lightness; the negative side includes highlight saturation, highlight warmth, highlight-shadow Lab b separation, Lab L-b correlation, and halation red-excess variables. Therefore, PC1 is best read as a cool/brighter blue-green feature direction on the high side and a warm-highlight/halation direction on the low side. PC2 was labeled the same way: its positive side includes weighted palette Lab distance, midtone warmth, RGB red mean, Lab chroma, and red-orange hue share, while its negative side includes teal-cyan hue share, blue hue share, shadow rolloff, and hue mean. Therefore, PC2 is mostly a palette-warmth/chroma direction opposed to teal-blue shadow structure.")
    add_figure(doc, "v2_01_pca_by_model.png", "Figure 2. PCA of the 195 standardized image features by model. Data: 288 images represented as 195 standardized image features. Each point is one generated image represented by PC1 and PC2 scores. After z-scoring all features, each PCA eigenvector was ranked by absolute loading magnitude; the largest positive and negative loadings were used to label the axes. High PC1 emphasizes blue-green brightness, low PC1 emphasizes warm highlights, and high PC2 emphasizes palette warmth/chroma. See Section 4.3 for the loading table. The ellipses show 80 percent covariance regions, not tiny confidence intervals around group means.", width_inches=6.35)
    add_figure(doc, "v2_02_pca_by_lighting.png", "Figure 3. PCA of the 195 standardized image features by lighting. Data: 288 images represented as 195 standardized image features. The same PC axes are used as in Figure 2. Warm practical and cool ambient conditions separate most clearly along feature directions tied to warmth, highlight color, local Lab structure, and teal-blue shadow behavior. The ellipses show 80 percent covariance regions.", width_inches=6.35)
    add_table(
        doc,
        ["Term", "Partial R²", "Pseudo F", "Permutation p", "FDR p"],
        permanova_rows(),
        [1.6, 0.85, 0.85, 0.95, 0.8],
        font_size=7.8,
        aligns=["left", "center", "center", "center", "center"],
    )
    add_figure(doc, "v2_05_blocked_permanova_partial_r2.png", "Figure 4. Blocked PERMANOVA on Euclidean distances among PC1-PC10 scores. The input data are 195 standardized image-feature measurements per generated image. PCA reduces those measurements to ten unitless component scores, and PERMANOVA partitions the distance variation by model, filmstock, lighting, scan, and interactions.", width_inches=6.35)
    add_para(doc, "The blocked PERMANOVA confirms the visual PCA pattern. The response being tested is the PC1-PC10 representation of the standardized image-feature matrix, so the partial R² values describe how much multivariate feature-space variation is associated with each design term. A partial R² of 0.118 for lighting means that lighting accounts for about 11.8 percent of the modeled PC1-PC10 distance variation after the other terms are included. Lighting is the largest factor, p = .001. Model-by-lighting is also significant, partial R² = 0.039, p = .001. Model alone is significant in multivariate feature space, partial R² = 0.037, p = .001, but this does not mean one model is universally more responsive. It means the generated images contain measurable model-specific feature structure.")

    doc.add_page_break()
    add_heading(doc, "5.2 Matched Prompt-Response Distances", 2)
    add_para(doc, "The all-feature responsiveness distance summarizes how far an image moves when one prompt factor changes. This gives a useful primary response variable, but it also compresses many types of visual movement into one number.")
    add_table(
        doc,
        ["Prompt effect", "ChatGPT mean", "Grok mean", "Grok minus ChatGPT", "95% CI", "p"],
        primary_distance_rows(),
        [1.8, 0.9, 0.9, 1.15, 1.4, 0.65],
        font_size=7.8,
        aligns=["left", "center", "center", "center", "center", "center"],
    )
    add_figure(doc, "v2_07_responsiveness_distance_bars_ci.png", "Figure 5. Prompt responsiveness distances by prompt factor, model, and feature family. Each bar is the mean matched-pair Euclidean distance after changing exactly one prompt factor while holding the other design variables fixed. Distances are computed across z-scored image features, so they summarize multivariate movement rather than raw pixel change. Lighting has the largest overall movement.", width_inches=6.3)
    add_para(doc, "The pooled paired tests show that total filmstock responsiveness is nearly identical across models: Grok minus ChatGPT = -0.035, p = .919. Scan responsiveness is also not significantly different: -0.383, p = .271. The significant difference appears for lighting: Grok minus ChatGPT = -1.429, p < .001, meaning ChatGPT moved more on the all-feature lighting distance.")
    add_figure(doc, "v2_09_paired_lines_lighting_distance.png", "Figure 6. Matched lighting responsiveness pairs. Each line compares models for the same run, scene, filmstock, scan condition, and replicate. Most matched conditions slope downward from ChatGPT to Grok, showing the model-by-lighting difference.", width_inches=5.6)
    add_table(
        doc,
        ["Prompt effect", "Mean difference", "Hierarchical bootstrap 95% CI", "Bootstrap p"],
        bootstrap_rows(),
        [2.0, 1.05, 1.9, 0.85],
        font_size=7.8,
        aligns=["left", "center", "center", "center"],
    )
    add_figure(doc, "v2_10_hierarchical_bootstrap_forest.png", "Figure 7. Hierarchical bootstrap CIs for Grok minus ChatGPT on all-feature responsiveness. The bootstrap resamples run blocks first and then matched contrasts inside selected runs, so the intervals preserve the three-run replicate structure. Only lighting excludes zero.", width_inches=6.2)

    add_heading(doc, "5.3 Lab, Split-Toning, and Halation Lighting Effects", 2)
    add_para(doc, "The targeted feature analysis explains what the lighting result means visually. Warm practical lighting changes not just average color, but the relationship between lightness and color across local regions, and the separation between highlights and shadows.")
    add_figure(doc, "v2_11_targeted_feature_heatmap_curated.png", "Figure 8. Curated targeted feature effects. Cell values are mean high-minus-low feature changes after each feature is divided by its sample standard deviation. Lighting dominates split-toning, warmth, Lab yellow-blue, and halation-related measures, so the statistical result maps onto identifiable photographic feature families rather than only numerical separation.", width_inches=5.7)
    add_figure(doc, "v2_13_top_lighting_screening_features.png", "Figure 9. Top model-differentiating lighting features from the all-feature screen. The x-axis is the absolute Grok-versus-ChatGPT difference in SD-scaled lighting effects for each feature. Local Lab covariance and split-toning features carry the strongest model gap.", width_inches=6.25)
    add_para(doc, "In the curated feature heatmap, ChatGPT shows larger lighting effects on split-tone Lab distance, highlight-shadow Lab b separation, warmth, Lab yellow-blue, and halation-related variables. The exploratory screen across all 195 features points in the same direction: the largest model gaps for lighting are local L-b covariance, local L-b correlation, split-tone distance, highlight saturation, and local a-b covariance. These are all interpretable photographic features rather than arbitrary pixel statistics.")

    doc.add_page_break()
    add_heading(doc, "5.4 CineStill-vs-Portra and Pushed-vs-Clean Scan Effects", 2)
    add_para(doc, "Filmstock and scan do not create the same dramatic overall movement as lighting, but they are not absent. In the feature screen, CineStill versus Portra is most associated with Fourier texture power, gradient magnitude, edge density, local color correlation, and palette chroma. Pushed scan versus clean scan appears in lower-tail luminance, palette distances, contrast ratios, shadow saturation, and texture measures.")
    add_table(
        doc,
        ["Feature", "Family", "Model gap", "FDR p"],
        top_screening_rows("filmstock_cinestill_minus_portra", 8),
        [2.45, 1.25, 0.9, 0.75],
        font_size=7.5,
        aligns=["left", "left", "center", "center"],
    )
    add_table(
        doc,
        ["Feature", "Family", "Model gap", "FDR p"],
        top_screening_rows("scan_pushed_minus_clean", 8),
        [2.45, 1.25, 0.9, 0.75],
        font_size=7.5,
        aligns=["left", "left", "center", "center"],
    )

    doc.add_page_break()
    add_heading(doc, "5.5 OLS, WLS, Huber IRLS, Kruskal-Wallis, and Friedman Results", 2)
    add_para(doc, "The regression analysis models all-feature responsiveness distance directly. The model-by-lighting coefficient is negative and significant in OLS with HC3 robust standard errors, meaning Grok's lighting response distance is lower than ChatGPT's after controlling for prompt factor, scene, run block, and replicate. WLS and Huber IRLS are included as sensitivity analyses for heteroskedasticity and influential outputs.")
    add_table(
        doc,
        ["Term", "OLS HC3 β (p)", "Two-stage WLS β (p)", "Huber IRLS β (p)"],
        regression_rows(),
        [1.4, 1.35, 1.35, 1.35],
        font_size=7.4,
    )
    add_para(doc, "The Kruskal-Wallis and Friedman results give the same conclusion: prompt effect type matters. Filmstock, lighting, and scan distances are not interchangeable categories. This matters because generated-image features do not necessarily satisfy clean normal-error assumptions.")

    add_heading(doc, "5.6 Leave-Run-Out PLS/Ridge/LASSO and Feature Variance Shares", 2)
    add_para(doc, "The predictive classification analyses ask a different question: can the extracted image features recover known experimental labels under leave-run-out validation? The answer is yes, especially for model and lighting. This does not replace inference, but it shows that the feature space contains strong and repeatable signal.")
    add_table(
        doc,
        ["Target", "Best PLS components", "AUC", "Accuracy", "RMSE"],
        pls_rows(),
        [1.2, 1.0, 0.8, 0.8, 0.8],
        font_size=7.8,
        aligns=["left", "center", "center", "center", "center"],
    )
    add_figure(doc, "v2_15_pls_vip_top_features.png", "Figure 10. PLS-DA VIP features by prediction target. VIP is a unitless predictor-importance score; values above 1 are usually treated as especially influential in the latent-variable classifier. Model and lighting labels are especially recoverable from the extracted features.", width_inches=6.4)
    add_table(
        doc,
        ["Target", "Method", "AUC", "Accuracy", "Median C"],
        regularized_rows(),
        [1.15, 1.3, 0.75, 0.75, 0.75],
        font_size=7.5,
        aligns=["left", "left", "center", "center", "center"],
    )
    doc.add_page_break()
    add_para(doc, "Feature-level variance decomposition is also encouraging, but its main statistical message is similarity rather than model separation. Prompt condition explains about 52.0 percent of ChatGPT feature variance and 55.0 percent of Grok feature variance. Pure run-block variance is tiny: about 0.8 percent for ChatGPT and 0.6 percent for Grok. Within-run replicate variation accounts for about 27.6 percent of ChatGPT feature variance and 26.1 percent of Grok feature variance, which justifies the extra runs and replicates.")
    add_table(
        doc,
        ["Model", "Prompt condition", "Run block", "Condition × run", "Within-run replicate"],
        variance_component_rows(),
        [1.3, 1.0, 0.9, 1.1, 1.2],
        font_size=7.8,
        aligns=["left", "center", "center", "center", "center"],
    )
    add_figure(doc, "v2_17_feature_variance_components.png", "Figure 11. Feature-level variance components. The top panel shows the mean share of total standardized feature variance across 195 features. The bottom panel shows Grok minus ChatGPT in percentage points. Both models have a similar variance structure: prompt condition accounts for the largest share, run-block variation is small, and within-run replicate variation remains measurable.", width_inches=6.15)
    add_figure(doc, "v2_18_seed_to_seed_variation.png", "Figure 12. Seed-to-seed variation within exact prompt conditions. The response is the distance between repeated generations for the same exact prompt within the same run. Within-run replicate distance is measurable but smaller than the prompt-condition signal.", width_inches=5.3)

    add_heading(doc, "5.7 Filmstock-by-Lighting-by-Scan Interaction Contrasts", 2)
    add_para(doc, "The factorial design also permits interaction contrasts. These ask whether prompt factors behave additively or whether the effect of one factor changes depending on another factor. For example, a filmstock-by-lighting interaction for one feature is computed as (CineStill warm - Portra warm) - (CineStill cool - Portra cool), then divided by that feature's sample standard deviation. The interaction heatmap suggests nonadditive factor behavior, especially in the three-way filmstock-by-lighting-by-scan contrast.")
    add_figure(doc, "v2_23_existing_interaction_heatmap.png", "Figure 13. Mean absolute factor interaction contrasts. Each cell averages absolute SD-scaled interaction contrasts across features and matched design cells. Larger values mean that the combined factors depart more from a simple additive pattern.", width_inches=5.4)

    doc.add_page_break()
    add_heading(doc, "6. Statistical Interpretation for Filmstock and Lighting Control", 1)
    add_para(doc, "The main interpretation is comparative rather than a single ranking of the two models. The evidence does not support a simple claim that ChatGPT Image 2 is better than Grok Imagine or that Grok Imagine is better than ChatGPT Image 2. It supports a more useful claim: photographic and cinematic responsiveness has dimensions. A model can be highly responsive to lighting but less distinctive on filmstock. A model can show texture movement for a pushed scan while showing only modest changes in halation. A single overall aesthetic score would flatten those distinctions.")
    add_para(doc, "The lighting result is the most robust finding. It appears in PCA, blocked PERMANOVA, paired distance tests, hierarchical bootstrap inference, feature-level screening, PLS-DA, and regularized classification. It is also photographically interpretable: the relevant features are warmth, Lab yellow-blue, split-tone distance, highlight-shadow color separation, and local Lab covariance. These are production-relevant descriptors for cinematography and color correction, not merely abstract statistical variables.")
    add_para(doc, "The model-by-lighting interaction is especially important statistically. If the only question were whether one model had a larger average response distance, then the conclusion would be too broad. The interaction gives a more specific result: the lighting factor does not move both models through feature space in exactly the same way. For a professional user, the relevant question is not only whether an image changes, but whether it changes along the intended visual dimension: warmer practicals, deeper local shadows, visible highlight color, tungsten/neon separation, and coherent color covariance across the frame.")
    add_para(doc, "Filmstock and scan should be interpreted with narrower effect sizes. They are detectable, but their effects are subtler and more feature-specific than lighting. This distinction is informative: the benchmark can separate prompt factors that visibly dominate the output from factors that require targeted metrics. It also warns against overinterpreting a non-significant overall distance. A filmstock factor might not greatly increase total Euclidean movement, yet still change the image in targeted feature families such as halation, palette separation, high-frequency texture, or shadow density.")
    add_para(doc, "The variance results also inform larger benchmark design. Pure run-block variation is small, which suggests that the three batches were not highly inconsistent. But within-run replicate variation is still substantial. In practical terms, two images from the same exact prompt can differ enough that a benchmark should not rely on one image per condition. The blocked replicate design gives a way to evaluate closed models responsibly: it cannot reveal the hidden sampler or true seed, but it can separate prompt-condition signal from repeated-generation variability.")
    add_para(doc, "There is also a methodological lesson about collinearity. Many photographic features are naturally correlated: warm highlights, Lab b, red-orange hue share, and split-tone distance are not independent artistic phenomena. PCA, PLS, ridge, and LASSO are helpful because they treat the feature space as structured rather than pretending every extracted variable is a separate clean measurement. In this study, that correlation structure is part of the image behavior being measured, not just a modeling nuisance.")
    add_heading(doc, "6.1 Limitations and Threats to Validity", 2)
    add_para(doc, "The study's scope is intentionally limited. First, both evaluated systems are closed cloud models, so seeds, samplers, internal model revisions, and some generation settings are not observable. Second, there are only three complete run blocks; therefore, block-level bootstrap intervals should be interpreted as sensitivity evidence from repeated batches rather than as large-sample asymptotic guarantees. Third, the feature set is handcrafted and should be understood as a proxy for photographic behavior, not a complete representation of human visual judgment or true chemical film response. Fourth, the named film stocks are prompt conditions, not ground-truth film simulations verified against reference photographs. Fifth, no expert human-rating panel was included, so the study measures statistical responsiveness rather than final aesthetic preference.")
    add_para(doc, "The industry relevance follows from that same point. Existing image benchmarks are highly useful when the question is whether a model follows object, counting, spatial, text, or compositional instructions. Creative work adds another layer: does the model understand the control language used by photographers, cinematographers, colorists, and art directors? In a production workflow, the value of an image model is partly its ability to respond predictably to terms like tungsten practical, pushed development, Frontier scan, halation, muted skin tone, or cool ambient window light. A benchmark in this style could inform model cards, creative-tool evaluation, product positioning, and procurement decisions for production workflows.")

    add_heading(doc, "7. Conclusion: Toward a Larger Filmmaker's Benchmark", 1)
    add_para(doc, "This study shows that photographic prompt responsiveness can be studied statistically. The benchmark generated 288 images, extracted 195 image features, and constructed blocked prompt-response distances as the primary response variable. The main inference used paired model comparisons, hierarchical bootstrap sensitivity analysis, PCA, blocked MANOVA, and blocked PERMANOVA, while regression, WLS, Huber robust regression, Kruskal-Wallis and Friedman tests, PLS, LASSO/ridge classification, variance decomposition, and robust outlier diagnostics provided additional checks on the same design.")
    add_para(doc, "The main empirical claim is nuanced: lighting is the strongest artistic control in this design, the model-by-lighting interaction is robust, and filmstock/scan factors are measurable but subtler. The larger conceptual claim is that creative users need benchmarks that measure controllability, not only correctness. A model that can render a scene correctly may still fail as a creative instrument if it cannot respond to the aesthetics of filmstock, lighting, texture, color separation, and postproduction treatment.")
    add_para(doc, "Filmstock language remains relevant even when images are produced computationally. Creative technologies inherit older visual vocabularies: photography borrowed from painting, early cinema borrowed from theater, and digital photography continues to borrow from film. AI image tools are likely to inherit and reinterpret photographic language in the same way. Prompts based on Kodak Portra, CineStill, pushed development, halation, lifted blacks, tinted shadows, tungsten balance, or lab scans are not merely retro labels; they name recognizable visual behaviors that many contemporary creators still seek.")
    add_para(doc, "Open-source image generation makes this concrete. Fine-tunes and lightweight adapters can steer a base model toward a particular visual domain, so a filmic look can become a calibrated image-generation pipeline rather than a vague aesthetic preference. If AI imagery is sometimes criticized for lacking aesthetic care, one constructive response is to build benchmarks that make aesthetic control explicit, measurable, and teachable.")
    add_para(doc, "Future work should expand the filmmaker's benchmark with more models, scenes, and repeats; expert ratings from photographers and art directors; reference photographs from named film stocks; open-model comparisons with controlled seeds; and additional variables such as lens focal length, camera height, staging density, production design, wardrobe palette, time of day, and color grade. The goal is to evaluate generative models as tools for artistic control, not only as systems for plausible image completion.")

    add_heading(doc, "Appendix A: Supplementary PCA, Distance, Feature, and Image Figures", 1)
    add_para(doc, "The following figures provide additional visual diagnostics for the multivariate analysis, responsiveness distances, feature screening, robustness diagnostics, and example images.")
    appendix_figures = [
        ("v2_03_pca_by_filmstock.png", "Appendix Figure A1. PCA by filmstock. Filmstock is detectable but less visually separated than lighting.", 6.1),
        ("v2_04_pca_by_run_block.png", "Appendix Figure A2. PCA by run block. The run blocks do not erase the main prompt structure.", 6.5),
        ("v2_06_manova_permanova_comparison.png", "Appendix Figure A3. MANOVA and PERMANOVA comparison. Both multivariate views place lighting at the top.", 6.3),
        ("v2_08_distance_box_strip_by_effect.png", "Appendix Figure A4. Box and strip plots for responsiveness distances.", 6.4),
        ("v2_12_all_feature_screening_volcano.png", "Appendix Figure A5. Exploratory all-feature screening across 195 features.", 6.3),
        ("v2_14_top_filmstock_screening_features.png", "Appendix Figure A6. Top filmstock model-differentiating features.", 6.2),
        ("v2_16_regularized_coefficients_model_lighting.png", "Appendix Figure A7. Ridge/LASSO-style coefficient diagnostics.", 6.2),
        ("v2_19_run_block_variation.png", "Appendix Figure A8. Run-block variation in responsiveness by prompt factor.", 5.8),
        ("v2_20_cross_run_reliable_features.png", "Appendix Figure A9. Feature effects with stable direction across all three runs.", 6.3),
        ("v2_21_robust_outlier_diagnostics.png", "Appendix Figure A10. Robust Mahalanobis outlier diagnostics.", 6.2),
        ("v2_24_existing_feature_correlation_heatmap.png", "Appendix Figure A11. Correlation matrix for selected image features. Pearson correlations were computed across the 288 generated images for a representative subset of color, split-tone, halation, texture, edge, palette, and local covariance features. The visible correlation blocks show that several feature columns measure related photographic responses rather than independent evidence, which motivates PCA, PLS, ridge/LASSO, and multivariate inference in the main analysis.", 5.9),
        ("v2_25_existing_nonparametric_ordering.png", "Appendix Figure A12. Nonparametric effect ordering plot.", 5.8),
        ("v2_27_contact_sheet_apartment.png", "Appendix Figure A13. Apartment editorial contact sheet.", 6.7),
        ("v2_28_contact_sheet_backstage.png", "Appendix Figure A14. Backstage musician contact sheet.", 6.7),
    ]
    for filename, caption, width in appendix_figures:
        add_figure(doc, filename, caption, width)

    doc.add_page_break()
    add_heading(doc, "Appendix B: Prompt Blocks", 1)
    add_table(
        doc,
        ["Prompt component", "Levels"],
        [
            ["Film stock", "Kodak Portra 400, 35mm color negative film, natural skin tones, soft pastel color response, gentle highlight rolloff; CineStill 800T, 35mm tungsten-balanced color negative film, cinematic color separation, pronounced red halation around bright highlights when present"],
            ["Lighting", "Cool diffuse ambient light, soft open shadows, no dominant warm practical light source; warm tungsten practical lights with subtle neon spill, visible small bright light sources in the frame, deeper localized shadows"],
            ["Development-scan", "Rated at box speed, clean neutral Frontier lab scan, restrained contrast, fine natural grain; pushed one stop in development, stronger contrast, denser shadows, coarser visible grain, slight color shift"],
            ["Prompt shell", "Square photorealistic editorial 35mm photograph of [scene]. Natural human proportions, realistic skin texture, art-directed but not glossy. Camera at eye level, 50mm lens, medium shot, subject clearly visible. [film]. [lighting]. [scan]. No text, no logos, no watermark, no artificial digital smoothness."],
        ],
        [1.2, 4.9],
        font_size=7.2,
    )

    doc.add_page_break()
    add_heading(doc, "Appendix C: Analysis Algorithm Outline", 1)
    add_numbered_item(doc, "Generate three complete balanced image batches: model × scene × filmstock × lighting × scan × replicate.")
    add_numbered_item(doc, "Validate every PNG: file exists, readable, RGB mode, square, 1024 × 1024 pixels.")
    add_numbered_item(doc, "Extract image features: color, Lab covariance, hue, split-toning, tonal curve, Fourier texture, palette structure, edge structure, and halation proxies.")
    add_numbered_item(doc, "Standardize feature columns.")
    add_equation(doc, "eq_appendix_standardization.png", r"z_{ij}=\frac{x_{ij}-\bar{x}_{j}}{s_{j}}")
    add_numbered_item(doc, "Construct blocked prompt-response distances.")
    add_equation(doc, "eq_appendix_distance.png", r"D_i=\sqrt{\sum_{j\in S}\left(z_{\mathrm{high},ij}-z_{\mathrm{low},ij}\right)^2}")
    add_numbered_item(doc, "Analyze with paired tests, hierarchical bootstrap, PCA, MANOVA, PERMANOVA, OLS HC3, WLS, Huber IRLS, nonparametric tests, PLS, LASSO/ridge, variance components, and robust outlier diagnostics.")

    doc.add_page_break()
    add_heading(doc, "References", 1)
    refs = [
        "Betker, J., Goh, G., Jing, L., Brooks, T., Wang, J., Li, L., Ouyang, L., Zhuang, J., Lee, J., Guo, Y., Manassra, W., Dhariwal, P., Chu, C., Jiao, Y., & Ramesh, A. (2023). Improving image generation with better captions. OpenAI. https://cdn.openai.com/papers/dall-e-3.pdf",
        "Butt, M. A., Gomez-Villa, A., Wu, T., Vazquez-Corral, J., van de Weijer, J., & Wang, K. (2025). GenColorBench: A color evaluation benchmark for text-to-image generation models. arXiv. https://arxiv.org/abs/2510.20586",
        "Ghosh, D., Hajishirzi, H., & Schmidt, L. (2023). GenEval: An object-focused framework for evaluating text-to-image alignment. arXiv. https://arxiv.org/abs/2310.11513",
        "Hu, Y., Liu, B., Kasai, J., Wang, Y., Ostendorf, M., Krishna, R., & Smith, N. A. (2023). TIFA: Accurate and interpretable text-to-image faithfulness evaluation with question answering. arXiv. https://arxiv.org/abs/2303.11897",
        "Huang, K., Duan, C., Sun, K., Xie, E., Li, Z., & Liu, X. (2023). T2I-CompBench++: An enhanced and comprehensive benchmark for compositional text-to-image generation. arXiv. https://arxiv.org/abs/2307.06350",
        "Lee, T., Yasunaga, M., Meng, C., Mai, Y., Park, J. S., Gupta, A., Zhang, Y., Narayanan, D., Teufel, H. B., Bellagente, M., Kang, M., Park, T., Leskovec, J., Zhu, J.-Y., Fei-Fei, L., Wu, J., Ermon, S., & Liang, P. (2023). Holistic evaluation of text-to-image models. arXiv. https://arxiv.org/abs/2311.04287",
        "Li, B., Lin, Z., Pathak, D., Li, J., Fei, Y., Wu, K., Ling, T., Xia, X., Zhang, P., Neubig, G., & Ramanan, D. (2024). GenAI-Bench: Evaluating and improving compositional text-to-visual generation. arXiv. https://arxiv.org/abs/2406.13743",
        "Lin, Z., Pathak, D., Li, B., Li, J., Xia, X., Neubig, G., Zhang, P., & Ramanan, D. (2024). Evaluating text-to-visual generation with image-to-text generation. arXiv. https://arxiv.org/abs/2404.01291",
        "Ruan, C., Xiao, Y., Hou, Y., Hu, G., & Zeng, W. (2026). ColorConceptBench: A benchmark for probabilistic color-concept understanding in text-to-image models. arXiv. https://arxiv.org/abs/2601.16836",
        "Saharia, C., Chan, W., Saxena, S., Li, L., Whang, J., Denton, E., Ghasemipour, S. K. S., Ayan, B. K., Mahdavi, S. S., Lopes, R. G., Salimans, T., Ho, J., Fleet, D. J., & Norouzi, M. (2022). Photorealistic text-to-image diffusion models with deep language understanding. arXiv. https://arxiv.org/abs/2205.11487",
        "Wu, X., Hao, Y., Sun, K., Chen, Y., Zhu, F., Zhao, R., & Li, H. (2023). Human Preference Score v2: A solid benchmark for evaluating human preferences of text-to-image synthesis. arXiv. https://arxiv.org/abs/2306.09341",
        "Yu, J., Xu, Y., Koh, J. Y., Luong, T., Baid, G., Wang, Z., Vasudevan, V., Ku, A., Yang, Y., Ayan, B. K., Hutchinson, B., Han, W., Parekh, Z., Li, X., Zhang, H., Baldridge, J., & Wu, Y. (2022). Scaling autoregressive models for content-rich text-to-image generation. arXiv. https://arxiv.org/abs/2206.10789",
    ]
    for ref in refs:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.5)
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.space_after = Pt(3)
        p.add_run(ref)

    OUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT_DOCX)


if __name__ == "__main__":
    build_doc()
    print(OUT_DOCX)
